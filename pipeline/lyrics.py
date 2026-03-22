import subprocess
from docx import Document

def clean_lyrics(text):
    return "\n".join([l.strip() for l in text.split("\n") if len(l.strip()) > 2])


def extract_lyrics(vocal_path):
    subprocess.run(
        f'whisper "{vocal_path}" --model large-v3 --language pt --output_format txt',
        shell=True,
        check=True
    )
    return f"{vocal_path}.txt"


from pipeline.versioning import version_lyrics

def create_docx(text_path):
    with open(text_path, "r", encoding="utf-8") as f:
        pt = clean_lyrics(f.read())

    en = version_lyrics(pt)

    doc = Document()

    doc.add_heading("Portuguese", 0)
    doc.add_paragraph(pt)

    doc.add_heading("English Version (EDIT REQUIRED)", 0)
    doc.add_paragraph(en)

    path = "lyrics/versioned.docx"
    doc.save(path)

    return path